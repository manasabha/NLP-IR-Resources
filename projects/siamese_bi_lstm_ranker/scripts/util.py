import csv
import ast
import numpy as np
import os, boto3, botocore
import json
from difflib import SequenceMatcher
import torch
import torch.nn as nn
import string, re
from torch.nn.utils.rnn import pad_sequence
from scripts.data_processing import TripleDataset, TripleDatasetNoShuffle

def create_all_data_paths():
    if not os.path.exists('data'):
        os.makedirs('data')

    if not os.path.exists('data/loss_graphs'):
        os.makedirs('data/loss_graphs')

    if not os.path.exists('data/common'):
        os.makedirs('data/common')

    if not os.path.exists('models'):
        os.makedirs('models')

    if not os.path.exists('models/Q'):
        os.makedirs('models/Q')

    if not os.path.exists('models/P'):
        os.makedirs('models/P')

    if not os.path.exists('models/attn'):
        os.makedirs('models/attn')
create_all_data_paths()

def upload_file_to_s3(file, s3_filename, s3_bucket_name, s3_path_to_filename_folder=None):
    """ Uploads a file to s3.
    Args:
        file (str): Path to the local file
        s3_filename (str): The name of the uploade file on s3.
        s3_bucket_name (str): The name of the s3 bucket.
        s3_path_to_filename_folder (str): The path to the folder where you want to save your file on s3.
    Returns:
        None
    """

    ACCESS_ID = os.getenv('ACCESS_ID')
    ACCESS_KEY = os.getenv('ACCESS_KEY')

    s3 = boto3.resource('s3',
                        aws_access_key_id=ACCESS_ID,
                        aws_secret_access_key=ACCESS_KEY)
    if not s3.Bucket(s3_bucket_name) in s3.buckets.all():
        raise Exception('A bucket with the name {bucket} does not exist on s3.'.format(bucket=s3_bucket_name))

    if s3_path_to_filename_folder:
        s3_filename = os.path.join(s3_path_to_filename_folder, s3_filename)
    s3.meta.client.upload_file(file, s3_bucket_name, s3_filename)


def get_length_of_dataset(input_type):
    length = 0
    with open('data/{}/ind2word.csv'.format(input_type)) as f:
        reader = csv.reader(f, delimiter='\t')
        for row in reader:
            length+=1
    return length

def similar(a, b):
    return SequenceMatcher(None, a, b).ratio()

def clean_text(text):
    text = text.lower()
    content = ''.join([i if ord(i) < 128 else ' ' for i in text])
    content = content.replace('[','').replace(']','')
    table = str.maketrans({key: None for key in string.punctuation})
    content = content.strip().rstrip(" ?:!.,;")
    content = content.translate(table)
    return content

def clean_sentence(sent):
    return ' '.join([clean_text(word) for word in sent.split(' ')])

def is_there_highlight_first(passage):
    findall_first = re.findall(r'\[\[[\w+\W+]+\]\]',passage[:351])
    if findall_first:
        return True
    else:
        return False

def is_there_highlight_last(passage):
    findall_last = re.findall(r'\[\[[\w+\W+]+\]\]',passage[-351:])
    if findall_last:
        return True
    else:
        return False

def read_word2_ind(input_type):
    word2ind = {}
    with open('data/{}/word_to_ind.csv'.format(input_type)) as f:
        reader = csv.reader(f, delimiter='\t')
        for row in reader:
            word2ind[row[0]] = int(row[1])
    return word2ind

def get_DRMM_vecs_for_ind(word2ind,wv_from_text):
    ind2vec = {}
    for word, ind in word2ind.items():
        if word in wv_from_text:
            ind2vec[ind] = wv_from_text[word]
    return ind2vec


def load_id_to_lpo_data():
    with open('data/lpo_data.json') as f:
        lpo_data = json.load(f)
    return lpo_data

def get_lpo_data_point_for_questions(list_of_questions, lpo_data):
    data = []
    for que in list_of_questions:
        data.append(get_question_id_and_full_entry_by_question(que, lpo_data))
    return data

def get_question_id_and_full_entry_by_question(que_to_check, lpo_data):
    for _id, que in lpo_data.items():
        question = que['question']

        cleaned_question = clean_sentence(question)
        if que_to_check == cleaned_question:
            return int(_id), que

#TODO make the original return a dict and index for _id
def get_n_validation_question_ids(validation_dataloader, lpo_dataset, ind2word):
    import random
    random_batch = random.randint(0, len(validation_dataloader))
    for i, (q, _, _) in enumerate(validation_dataloader):
        if i< random_batch:
            continue
        question_words_list = return_word_for_ind_list_batch(q, ind2word)
        mongo_question_ids = [get_question_id_and_full_entry_by_question(que, lpo_dataset)[0] for que in question_words_list]
        break
    return mongo_question_ids

def get_id_for_raw_question_lpo(question, lpo_dataset):
    for _id,entry in lpo_dataset.items():
        if entry['question'] == question:
            return _id

def get_raw_qualities_for_question(question, lpo_dataset):
    _id = get_id_for_raw_question_lpo(question, lpo_dataset)
    entry = lpo_dataset[_id]
    qualities = [ans['qua'] for ans in entry['answers']]
    return qualities

def select_question_entries_for_question_ids(validation_dataloader, lpo_dataset, ind2word):
    question_ids = get_n_validation_question_ids(validation_dataloader, lpo_dataset, ind2word)
    data = []
    for _id in question_ids:
        datum = {}
        entry = lpo_dataset[str(_id)]
        datum['question'] = entry['question']
        datum['answers'] = [ans['ans'] for ans in entry['answers']]
        data.append(datum)
    return data

def select_question_entries(lpo_dataset, no_of_questions):
    data = []
    for i, (_, entry) in enumerate(lpo_dataset.items()):
        datum = {}
        if i+1==no_of_questions:
            return data
        datum['question'] = entry['question']
        datum['answers'] = [ans['ans'] for ans in entry['answers']]
        data.append(datum)
    return data


def download_model_from_s3(model_name):
    ACCESS_ID = os.getenv('ACCESS_ID')
    ACCESS_KEY = os.getenv('ACCESS_KEY')
    BUCKET_NAME = 'manasa-mscac-project'
    Q_PATH = 'models/Q/{}'.format(model_name)
    P_PATH = 'models/P/{}'.format(model_name)

    if os.path.isfile(Q_PATH) and os.path.isfile(P_PATH):
        return

    s3_Q_PATH = 'models/Q/{}'.format(model_name)
    s3_P_PATH = 'models/P/{}'.format(model_name)
    s3 = boto3.resource('s3',
                        aws_access_key_id=ACCESS_ID,
                        aws_secret_access_key=ACCESS_KEY)

    try:
        s3.Bucket(BUCKET_NAME).download_file(s3_Q_PATH, Q_PATH)
        s3.Bucket(BUCKET_NAME).download_file(s3_P_PATH, P_PATH)
    except botocore.exceptions.ClientError as e:
        if e.response['Error']['Code'] == "404":
            print("The object does not exist.")
        else:
            raise

def read_ind2word(input_type):

    ind2word = {}
    with open('data/{}/ind2word.csv'.format(input_type)) as f:
        reader = csv.reader(f, delimiter='\t')
        for row in reader:
            ind2word[int(row[0])] = row[1]
    return ind2word

def read_ind2vec(input_type):
    ind2vec = {}
    with open('data/{}/ind2vec.csv'.format(input_type)) as f:
        reader = csv.reader(f, delimiter='\t')
        for row in reader:
            list_vec = ast.literal_eval(row[1])
            ind2vec[int(row[0])] = np.array(list_vec)
    return ind2vec

def read_all_word_ind_vec_to_files(input_type):
    ind2vec = {}
    with open('data/{}/ind2vec.csv'.format(input_type)) as f:
        reader = csv.reader(f, delimiter='\t')
        for row in reader:
            list_vec = ast.literal_eval(row[1])
            ind2vec[int(row[0])] = np.array(list_vec)

    ind2word = {}
    with open('data/{}/ind2word.csv'.format(input_type)) as f:
        reader = csv.reader(f, delimiter='\t')
        for row in reader:
            ind2word[int(row[0])] = row[1]

    word2ind = {}
    with open('data/{}/word_to_ind.csv'.format(input_type)) as f:
        reader = csv.reader(f, delimiter='\t')
        for row in reader:
            word2ind[row[0]] = int(row[1])
    return ind2vec, ind2word, word2ind


def write_all_word_ind_vec_to_files(ind2vec, ind2word, word2ind, input_type):
    with open('data/{}/ind2vec.csv'.format(input_type), 'w') as f:
        writer = csv.writer(f, delimiter = '\t')
        for ind, vec in ind2vec.items():
            writer.writerow([ind, vec.tolist()])

    with open('data/{}/ind2word.csv'.format(input_type), 'w') as f:
        writer = csv.writer(f, delimiter = '\t')
        for ind, word in ind2word.items():
            writer.writerow([ind, word])

    with open('data/{}/word_to_ind.csv'.format(input_type), 'w') as f:
        writer = csv.writer(f, delimiter = '\t')
        for word,ind in word2ind.items():
            writer.writerow([word, ind])

def write_all_word_ind_vec_to_files_common(ind2vec, ind2word, word2ind, input_type='common'):
    with open('data/{}/ind2vec.csv'.format(input_type), 'w') as f:
        writer = csv.writer(f, delimiter = '\t')
        for ind, vec in ind2vec.items():
            writer.writerow([ind, vec.tolist()])

    with open('data/{}/ind2word.csv'.format(input_type), 'w') as f:
        writer = csv.writer(f, delimiter = '\t')
        for ind, word in ind2word.items():
            writer.writerow([ind, word])

    with open('data/{}/word_to_ind.csv'.format(input_type), 'w') as f:
        writer = csv.writer(f, delimiter = '\t')
        for word,ind in word2ind.items():
            writer.writerow([word, ind])

def write_rows_csv(file_name, data):
    with open(file_name, 'w') as f:
        writer = csv.writer(f, delimiter='\t')
        writer.writerows(data)

def return_ind_for_word_list(word_list, word2ind):
    inds = []
    missed = []
    for word in word_list.split(' '):
        try:
            inds.append(word2ind[clean_text(word)])
        except:
            missed.append(word)
    # missed_percentage = float(len(missed))/len(word_list.split(' '))
    # print('{} % of words missed total words {} missed words {}'.format(missed_percentage, len(word_list.split(' ')),
    #                                                                    len(missed)))
    # print('Missed words ', missed)
    return inds

def get_user_input(input_fields_with_default_values):
    from time import sleep
    new_value_dict = {}

    for field, def_val in input_fields_with_default_values.items():
        try:
            print('Please provide {} in 30 seconds! (Hit Ctrl-C to start)'.format(field))
            for i in range(0, 5):
                sleep(1)  # could use a backward counter to be preeety :)
            print('No input is given for {} setting to default value'.format(field))
            new_value_dict[field] = def_val
        except KeyboardInterrupt:
            field_val = input('Input {}:'.format(field))
            new_value_dict[field] = field_val
    return new_value_dict

def date_as_float():
    from datetime import datetime
    return ''.join(str(datetime.now()).replace('-', '').replace(':', '').replace('.', '').split(' '))[:12]


def return_word_for_ind_list_batch(ind_lists, ind2word):
    words = []
    for ind_list in ind_lists:
        words.append(return_word_for_ind_list(ind_list, ind2word))
    return words

def return_word_for_ind_list(ind_list, ind2word):
    if not isinstance(ind_list, list):
        ind_list = ind_list.numpy().tolist()
    return ' '.join([ind2word[ind] for ind in ind_list if ind])

def read_to_ind_files(file_name):
    data = []
    with open(file_name) as f:
        reader = csv.reader(f, delimiter='\t')
        for row in reader:
            list_vec = [int(r) for r in row]
            data.append(list_vec)
    return data

def read_files_and_return_dataset(device, split=0.7, is_common=False, is_shuffle=True):

    if is_common:
        path = 'data/common/'
    else:
        path = 'data/'
    print('Loading X_Q')
    X_Q = read_to_ind_files(path+'question_to_ids.csv')

    print('Loading X_pa1')
    X_pa1 = read_to_ind_files(path+'pa1_to_ids.csv')

    print('Loading X_pa2')
    X_pa2 = read_to_ind_files(path+'pa2_to_ids.csv')

    #Remove any empty entries - God knows why they are there
    X_Q = list(filter(None, X_Q))
    X_pa1 = list(filter(None, X_pa1))
    X_pa2 = list(filter(None, X_pa2))
    if is_shuffle:
        torch.manual_seed(2055)
        rand_inds = torch.randperm(len(X_Q)).to(device)
        triple_set_train = TripleDataset(X_Q, X_pa1, X_pa2, rand_inds, split=split, is_train=True)
        triple_set_val = TripleDataset(X_Q, X_pa1, X_pa2, rand_inds, split=split)
        torch.save(triple_set_train, 'data/train_dataset')
        torch.save(triple_set_val, 'data/val_dataset')
        del(rand_inds)
        torch.cuda.empty_cache()
    else:
        triple_set_train = TripleDatasetNoShuffle(X_Q, X_pa1, X_pa2, split=0.7, is_train=True)
        triple_set_val = TripleDatasetNoShuffle(X_Q, X_pa1, X_pa2, split=0.7)

    return triple_set_train, triple_set_val

def load_models(model_name, device, parameters, is_eval=False):
    download_model_from_s3(model_name)
    loaded_models = {
        'question_encoder': torch.load('models/Q/{}'.format(model_name), map_location=device),
        'passage_encoder': torch.load('models/P/{}'.format(model_name), map_location=device),
        'attention_model': None
    }
    if parameters['enable_attention'] and parameters['use_common_attention_model']:
        loaded_models['attn_model'] = torch.load('models/attn/{}'.format(model_name), map_location=device)

    if is_eval:
        for _, model in loaded_models.items():
            if not (model is None):
                model.eval()
    return loaded_models

def instantiate_and_load_models_new_flags(model_name, device, use_question_context,
                                          return_context, parameters, is_eval=True):
    from scripts.models import QuestionEncoder, PassageEncoder, Attention
    loaded_models = load_models(model_name, device, parameters, is_eval=is_eval)

    models = {'question_encoder': QuestionEncoder(return_context=return_context, **parameters).to(device),
              'passage_encoder': PassageEncoder(use_question_context=use_question_context, **parameters).to(device),
              'attention_model': None}

    if not validate_user_choices(parameters):
        print('Instantiate models Follow the instructions provided and rerun the script')

    if parameters['enable_attention'] and parameters['use_common_attention_model']:
        models['attn_model'] = Attention(parameters['attn_units']).to(device)

    for model_type, model in models.items():
        if not (model is None):
            model.load_state_dict(loaded_models[model_type].state_dict())
            if is_eval:
                model.eval()
    return models

def single_set_return_attn(question, list_of_passages, q_encoder, p_encoder, word2ind, device, attn_model=None):
    que2id = torch.LongTensor(return_ind_for_word_list(question, word2ind)).unsqueeze(0)
    list_indexed_passages = [torch.LongTensor(return_ind_for_word_list(passage, word2ind)).unsqueeze(0) for passage in list_of_passages]
    batch_to_data_q = {}
    batch_to_data_pa = {}
    p_attn_list = []
    q_enc_output = q_encoder(que2id, batch_to_data_q, 0, device, attn_model)
    for i, passage in enumerate(list_indexed_passages):
        p_enc_output = p_encoder(passage, batch_to_data_pa, i, device, q_enc_output['context_vector'], attn_model)
        p_attn_list.append(p_enc_output['attn_weights'].detach().numpy().tolist())
    return q_enc_output['attn_weights'].detach().numpy().tolist(), p_attn_list


def rank_passages_for_single_set_return_scores(question, list_of_passages,
                                               q_encoder, p_encoder,
                                               word2ind, device, attn_model=None):
    que2id = torch.LongTensor(return_ind_for_word_list(question, word2ind)).unsqueeze(0).to(device)
    list_indexed_passages = [torch.LongTensor(return_ind_for_word_list(passage, word2ind)) for passage in list_of_passages]
    passage_ids = pad_sequence(list_indexed_passages, batch_first=True).to(device)
    batch_to_data_q = {}
    batch_to_data_pa = {}

    q_enc_output = q_encoder(que2id, batch_to_data_q, 0, device, attn_model)
    p_enc_output = p_encoder(passage_ids, batch_to_data_pa, 0, device,
                             q_enc_output['encoded_question_vector'], attn_model)
    cos = nn.CosineSimilarity(dim=1, eps=1e-6)
    cos_sim = cos(q_enc_output['encoded_question_vector'],
                  p_enc_output['encoded_passage_vector']).detach().numpy().tolist()
    return cos_sim

def validate_user_choices(parameters):
    if parameters['use_common_attention_model'] or parameters['save_weights']:
        print('common attn {} save_weig {}'.format(parameters['use_common_attention_model'], parameters['save_weights']))
        if not parameters['enable_attention']:
            print('To use common attention or to save attention weights, please set the enable attention flag')
            return False
    return True


def instantiate_models(parameters, device):
    from scripts.models import QuestionEncoder, PassageEncoder, Attention
    new_parameters = parameters.copy()
    use_question_context = new_parameters.pop('use_question_context')
    return_context = new_parameters.pop('return_context')
    models = {'q_enc': QuestionEncoder(return_context=return_context, **new_parameters).to(device),
              'p_enc': PassageEncoder(use_question_context=use_question_context, **new_parameters).to(device),
              'cos': nn.CosineSimilarity(dim=1, eps=1e-6),
              'attn_model': None}

    if not validate_user_choices(parameters):
        print('Instantiate models Follow the instructions provided and rerun the script')

    if parameters['enable_attention'] and parameters['use_common_attention_model']:
        models['attn_model'] = Attention(parameters['attn_units']).to(device)
    return models

def write_to_word_doc(out_path, word_list, attn, document):
    from docx.shared import Pt, RGBColor
    sorted_attn = sorted(attn)
    sorted_ind = [sorted_attn.index(att) for att in attn]
    color_ind = [255 - int(255 * (float(ind) / len(sorted_ind))) for ind in sorted_ind]
    color_ind = [0 if c_ind < 0 else c_ind for c_ind in color_ind]
    # color_ind = [int(255 * (float(ind) / len(sorted_ind))) + 25 for ind in sorted_ind]
    # color_ind = [255 if c_ind > 255 else c_ind for c_ind in color_ind]
    para = document.add_paragraph()
    for i, word in enumerate(word_list):
        run = para.add_run(' ' + word)
        font = run.font
        if max(sorted_ind)>30:
            font.size = Pt(sorted_ind[i])
        else:
            font.size = Pt(sorted_ind[i]+10)
        font.color.rgb = RGBColor(color_ind[i], color_ind[i], color_ind[i])

    document.save(out_path)

def write_lists_to_word_doc(out_path, list_of_word_lists, list_of_attn):
    from docx import Document
    document = Document()

    for word_list, attn in zip(list_of_word_lists, list_of_attn):
        write_to_word_doc(out_path, word_list, attn, document)

def np_collate(batch):
    device = torch.device('cuda' if torch.cuda.is_available() else 'cpu')
    que = [torch.from_numpy(np.array(entry[0])) for entry in batch]
    bp = [torch.from_numpy(np.array(entry[1])) for entry in batch]
    p = [torch.from_numpy(np.array(entry[2])) for entry in batch]
    padded_data = pad_sequence(que, batch_first=True)
    passed_bp = pad_sequence(bp, batch_first=True)
    padded_p = pad_sequence(p, batch_first=True)
    return [torch.LongTensor(padded_data).to(device), torch.LongTensor(passed_bp).to(device),
            torch.LongTensor(padded_p).to(device)]
